from cmd import Cmd
from glob import glob
from typing import IO, List
from docx import Document
from docx.text.paragraph import Paragraph
from time import sleep
from colorama import Fore, init
from collections import OrderedDict
from natsort import natsorted
from .exceptions import *
from .functions import log, choice
from .config import ftfpath

import os
import re
import importlib
import datetime
import math
import getpass
import bcrypt

init()

def loadcmd():
    for i in glob("libs\\ExternalCommands\\cmd_*.py"):
        plugin_file_name = i.replace("\\", ".")[0:-3]
        log("External Commands: " + plugin_file_name, "debug")
        importlib.import_module(plugin_file_name)

class FTFCmd(Cmd):
    intro = "欢迎使用《朝花夕拾协议》命令行\n"
    prompt = "[FTF Command Line] "
    doc_header = "发现了以下命令的帮助文档，输入help <命令>来查看其帮助文档:"
    undoc_header = "未发现以下命令的帮助文档，也许这些命令有显示自己帮助文档的参数，试试输入<命令> /?来查看它们:"
    misc_header = "其它帮助命令:"
    nohelp = "没有找到%s命令的帮助文档，也许这个命令有显示自己帮助文档的参数，试试输入<命令> /?来查看它们"

    def __init__(self, completekey: str = "tab", stdin: IO[str] | None = None, stdout: IO[str] | None = None) -> None:
        super().__init__(completekey, stdin, stdout)
        self.YEARS = natsorted(i for i in os.listdir(ftfpath) if os.path.isdir(os.path.join(ftfpath, i)))
        self.COLOR = Fore.LIGHTGREEN_EX
        self.POSITIVE_LEVELS = {"高", "中", "低"}
        self.NEGATIVE_LEVELS = {"严重", "中", "轻微"}
        self.POSITIVE_ASSESS = {"积极一", "积极二", "积极三"}
        self.NEGATIVE_ASSESS = {"消极一", "消极二", "消极三"}
        self.WEEKLY_JUDGMENT = {"糟糕的一周", "平平无奇的一周", "标准的一周", "杰出的一周", "优秀的一周"}
        self.MONTHLY_JUDGMENT = {"糟糕的一个月", "平平无奇的一个月", "标准的一个月", "杰出的一个月", "优秀的一个月"}
        self.YEARLY_JUDGMENT = {"糟糕的一年", "平平无奇的一年", "标准的一年", "杰出的一年", "优秀的一年"}

    def emptyline(self):
        return
    
    def preloop(self) -> None:
        os.system("cls")
        loadcmd()

    def default(self, line: str) -> None:
        os.system(line)

    def onecmd(self, line: str) -> bool:
        if not line == "" and not line.isspace():
            log(self.prompt + line, "info", logfile_only=True)
        try:
            callback = super().onecmd(line)
        except Exception as e:
            if isinstance(e, AdminMode) or isinstance(e, CommandLineExit):
                raise e
            print(f"{Fore.LIGHTRED_EX}{e.__class__.__name__}: {str(e)}")
            log(f"{e.__class__.__name__}: {str(e)}", "error", logfile_only=True)
            print(f"{Fore.LIGHTRED_EX}运行时发生错误，详细信息已被写入日志")
            log("运行时发生错误，详细信息已被写入日志", "exception", logfile_only=True)
            return False
        return callback

    def do_exit(self, args: str):
        """
        退出《朝花夕拾协议》命令行。

        语法：exit [/?]
            无参数  退出《朝花夕拾协议》命令行。
            /?      显示此帮助文档。
        """
        if args.split(" ")[0] == "/?":
            print(self.do_exit.__doc__)
            return
        raise CommandLineExit()
    
    def _find(self, document: str, keywords: set[str], count: int) -> int:
        if os.path.basename(document).startswith("~$"):
            return count
        doc = Document(document)
        document_friendly_name = document.split("\\")[-2] + "/" + document.split("\\")[-1].removesuffix(".docx")
        line = 0
        first_month_span = None
        second_month_span = None
        half_length = len(doc.paragraphs) // 2
        for paragraph in doc.paragraphs:
            line += 1
            if paragraph.text == "跨月" or paragraph.text == "跨年":
                if not first_month_span:
                    first_month_span = line
                elif not second_month_span:
                    second_month_span = line
                else:
                    print(f"在{document_friendly_name}中发现两个以上的跨月或跨年事件，请检查文档")
                    log(f"在{document_friendly_name}中发现两个以上的跨月或跨年事件，请检查文档", "warning", logfile_only=True)
                    return count
        line = 0
        for paragraph in doc.paragraphs:
            line += 1
            for keyword in keywords:
                if keyword.startswith("/") and keyword.endswith("/search"):
                    if re.search(keyword[1:-7], paragraph.text):
                        tag = None
                        if keyword != "跨月":
                            if first_month_span:
                                if line < first_month_span and first_month_span < half_length:
                                    tag = "上跨月"
                                    colored_tag = f"{Fore.LIGHTBLUE_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                    self.has_upper_part = True
                                elif line > first_month_span and first_month_span > half_length:
                                    tag = "下跨月"
                                    colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                    self.has_lower_part = True
                            if second_month_span and line > second_month_span:
                                tag = "下跨月"
                                colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                self.has_lower_part = True
                        if tag:
                            print(f"{count + 1}. {colored_tag}在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-6]}的项（从整个字符串中搜索匹配项） -> {paragraph.text}")
                            log(f"{count + 1}. [{tag}]在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-6]}的项（从整个字符串中搜索匹配项） -> {paragraph.text}", "info", logfile_only=True)
                        else:
                            print(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-6]}的项（从整个字符串中搜索匹配项） -> {paragraph.text}")
                            log(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-6]}的项（从整个字符串中搜索匹配项） -> {paragraph.text}", "info", logfile_only=True)
                        count += 1
                elif keyword.startswith("/") and keyword.endswith("/match"):
                    if re.match(keyword[1:-6], paragraph.text):
                        tag = None
                        if keyword != "跨月":
                            if first_month_span:
                                if line < first_month_span and first_month_span < half_length:
                                    tag = "上跨月"
                                    colored_tag = f"{Fore.LIGHTBLUE_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                    self.has_upper_part = True
                                elif line > first_month_span and first_month_span > half_length:
                                    tag = "下跨月"
                                    colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                    self.has_lower_part = True
                            if second_month_span and line > second_month_span:
                                tag = "下跨月"
                                colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                self.has_lower_part = True
                        if tag:
                            print(f"{count + 1}. {colored_tag}在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-5]}的项（从字符串开头匹配） -> {paragraph.text}")
                            log(f"{count + 1}. [{tag}]在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-5]}的项（从字符串开头匹配） -> {paragraph.text}", "info", logfile_only=True)
                        else:
                            print(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-5]}的项（从字符串开头匹配） -> {paragraph.text}")
                            log(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到符合正则表达式{keyword[0:-5]}的项（从字符串开头匹配） -> {paragraph.text}", "info", logfile_only=True)
                        count += 1
                elif "&" in keyword:
                    if all(i in paragraph.text for i in keyword.split("&")):
                        tag = None
                        if keyword != "跨月":
                            if first_month_span:
                                if line < first_month_span and first_month_span < half_length:
                                    tag = "上跨月"
                                    colored_tag = f"{Fore.LIGHTBLUE_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                    self.has_upper_part = True
                                elif line > first_month_span and first_month_span > half_length:
                                    tag = "下跨月"
                                    colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                    self.has_lower_part = True
                            if second_month_span and line > second_month_span:
                                tag = "下跨月"
                                colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{Fore.LIGHTGREEN_EX}"
                                self.has_lower_part = True
                        if tag:
                            print(f"{count + 1}. {colored_tag}在{document_friendly_name}第{line}个段落中找到同时包含关键字词“{','.join(keyword.split('&'))}”的项 -> {paragraph.text}")
                            log(f"{count + 1}. [{tag}]在{document_friendly_name}第{line}个段落中找到同时包含关键字词“{','.join(keyword.split('&'))}”的项 -> {paragraph.text}", "info", logfile_only=True)
                        else:
                            print(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到同时包含关键字词“{','.join(keyword.split('&'))}”的项 -> {paragraph.text}")
                            log(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到同时包含关键字词“{','.join(keyword.split('&'))}”的项 -> {paragraph.text}", "info", logfile_only=True)
                        count += 1
                else:
                    if keyword in paragraph.text:
                        tag = None
                        if keyword != "跨月":
                            if first_month_span:
                                if line < first_month_span and first_month_span < half_length:
                                    tag = "上跨月"
                                    colored_tag = f"{Fore.LIGHTBLUE_EX}[{tag}]{self.COLOR}"
                                    self.has_upper_part = True
                                elif line > first_month_span and first_month_span > half_length:
                                    tag = "下跨月"
                                    colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{self.COLOR}"
                                    self.has_lower_part = True
                            if second_month_span and line > second_month_span:
                                tag = "下跨月"
                                colored_tag = f"{Fore.LIGHTYELLOW_EX}[{tag}]{self.COLOR}"
                                self.has_lower_part = True
                        if tag:
                            print(f"{count + 1}. {colored_tag}在{document_friendly_name}第{line}个段落中找到关键字词: {keyword} -> {paragraph.text}")
                            log(f"{count + 1}. [{tag}]在{document_friendly_name}第{line}个段落中找到关键字词: {keyword} -> {paragraph.text}", "info", logfile_only=True)
                        else:
                            print(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到关键字词: {keyword} -> {paragraph.text}")
                            log(f"{count + 1}. 在{document_friendly_name}第{line}个段落中找到关键字词: {keyword} -> {paragraph.text}", "info", logfile_only=True)
                        count += 1
        return count

    def _check_month_span(self):
        if self.has_upper_part or self.has_lower_part:
            print("注意：查询涉及的文档中检测到跨月或跨年事件，请注意核实查询结果的事件所属月份")
            log("注意：查询涉及的文档中检测到跨月或跨年事件，请注意核实查询结果的事件所属月份", "info", logfile_only=True)
            if self.has_upper_part:
                print(f"{Fore.LIGHTBLUE_EX}[上跨月]{self.COLOR}表示该事件属于查询结果所示文档的上个月")
                log("[上跨月]表示该事件属于查询结果所示文档的上个月", "info", logfile_only=True)
            if self.has_lower_part:
                print(f"{Fore.LIGHTYELLOW_EX}[下跨月]{self.COLOR}表示该事件属于查询结果所示文档的下个月")
                log("[下跨月]表示该事件属于查询结果所示文档的下个月", "info", logfile_only=True)

    def do_find(self, args: str):
        """
        查找文档中的关键字词。

        语法：find <keywords>/<regex> in [year <years> | month <months> | only <document> | *] [/?]
            keywords        指定查找的关键字词。
                            使用空格分隔多个关键字词，表示查找包含多个关键字词的项。
                            使用“&”符号连接多个关键字词，表示查找的项中必须同时包含这些关键字词。
            regex           正则表达式。以“/”符号开头并以“/search”或“/match”结尾的字符串将被视为正则表达式。
                            正则表达式以“/search”结尾表示从整个字符串中搜索匹配项。
                            正则表达式以“/match”结尾表示从字符串开头匹配。
            year <years>    文档所属的年份，可填多个，使用空格分隔。
                            也可使用“~”表示范围，如“2023~2026”表示2023年到2026年（包含2023年和2026年）的事件记录文档。
            month <months>  文档所属的月份，可填多个，使用空格分隔。
                            也可使用“~”表示范围，如“1~3”表示1月到3月（包含1月和3月）的事件记录文档。
            only <document> 仅在指定的文档中查找关键字词。
                            若指定的文档格式为<year>/<month>，则表示在事件记录文档中查找，
                            如“2023/1”表示2023年1月的事件记录文档。
                            若指定的文档格式为<year>/annual_summary，则表示在年度总结中查找，
                            如2023/annual_summary表示2023年的年度总结。
            *               在所有文档中查找关键字词，包括年度总结。
            /?              显示此帮助文档。
        """
        if args.split(" ")[0] == "/?" or args == "":
            print(self.do_find.__doc__)
            return
        self.has_upper_part = False
        self.has_lower_part = False
        keywords = set(args.split(" in ")[0].split(" "))
        documents = list(OrderedDict.fromkeys(args.split(" in ")[1].split(" ")).keys())
        if documents[0] == "*":
            count = 0
            for year in self.YEARS:
                docments_path = os.path.join(ftfpath, year)
                for document in natsorted(glob(f"{docments_path}\\*.docx")):
                    count = self._find(document, keywords, count)
            self._check_month_span()
            print(f"在所有文档中共发现{count}个关键字词\n")
            log(f"在所有文档中共发现{count}个关键字词", "info", logfile_only=True)
            log("", "info", logfile_only=True)
        elif documents[0] == "year":
            years = documents[1:]
            deduplicated_year = set()
            for year in years:
                if "~" in year:
                    start_year, end_year = year.split("~")
                    if not start_year.isdigit() or not end_year.isdigit():
                        print(f"无效的年份范围: {year}")
                        log(f"无效的年份范围: {year}", "warning", logfile_only=True)
                        years.remove(year)
                        continue
                    start_year, end_year = int(start_year), int(end_year)
                    if start_year > end_year:
                        print(f"无效的年份范围: {year}（起始年份大于结束年份）")
                        log(f"无效的年份范围: {year}（起始年份大于结束年份）", "warning", logfile_only=True)
                        years.remove(year)
                        continue
                    deduplicated_year.update(str(i) for i in range(start_year, end_year + 1))
                else:
                    deduplicated_year.add(year)
            deduplicated_year = natsorted(deduplicated_year)
            count = 0
            for year in deduplicated_year:
                if year not in self.YEARS:
                    print(f"未找到{year}年的事件记录文档")
                    log(f"未找到{year}年的事件记录文档", "warning", logfile_only=True)
                    continue
                docments_path = os.path.join(ftfpath, year)
                for document in natsorted(glob(f"{docments_path}\\*.docx")):
                    count = self._find(document, keywords, count)
            self._check_month_span()
            print(f"在{', '.join(years)}这{len(deduplicated_year)}年的事件记录文档中共发现{count}个关键字词\n")
            log(f"在{', '.join(years)}这{len(deduplicated_year)}年的事件记录文档中共发现{count}个关键字词", "info", logfile_only=True)
            log("", "info", logfile_only=True)
        elif documents[0] == "month":
            months = documents[1:]
            deduplicated_month = set()
            for month in months:
                if "~" in month:
                    start_month, end_month = month.split("~")
                    if not start_month.isdigit() or not end_month.isdigit():
                        print(f"无效的月份范围: {month}")
                        log(f"无效的月份范围: {month}", "warning", logfile_only=True)
                        months.remove(month)
                        continue
                    start_month, end_month = int(start_month), int(end_month)
                    if start_month > end_month:
                        print(f"无效的月份范围: {month}（起始月份大于结束月份）")
                        log(f"无效的月份范围: {month}（起始月份大于结束月份）", "warning", logfile_only=True)
                        months.remove(month)
                        continue
                    if start_month < 1 or end_month > 12:
                        print(f"无效的月份范围: {month}（月份应在1到12之间）")
                        log(f"无效的月份范围: {month}（月份应在1到12之间）", "warning", logfile_only=True)
                        months.remove(month)
                        continue
                    deduplicated_month.update(str(i) for i in range(start_month, end_month + 1))
                else:
                    deduplicated_month.add(month)
            deduplicated_month = natsorted(deduplicated_month)
            count = 0
            for year in self.YEARS:
                for month in deduplicated_month:
                    if month not in [str(i) for i in range(1, 13)]:
                        print(f"无效的月份: {month}月")
                        log(f"无效的月份: {month}月", "warning", logfile_only=True)
                        continue
                    docments_path = os.path.join(ftfpath, year, f"{month}月.docx")
                    if not os.path.exists(docments_path):
                        print(f"未找到{year}年{month}月的事件记录文档")
                        log(f"未找到{year}年{month}月的事件记录文档", "warning", logfile_only=True)
                        continue
                    count = self._find(docments_path, keywords, count)
            self._check_month_span()
            print(f"在{', '.join(months)}月这{len(deduplicated_month)}个月的事件记录文档中共发现{count}个关键字词\n")
            log(f"在{', '.join(months)}月这{len(deduplicated_month)}个月的事件记录文档中共发现{count}个关键字词", "info", logfile_only=True)
            log("", "info", logfile_only=True)
        elif documents[0] == "only":
            count = 0
            normal_count = 0
            annual_summary_count = 0
            document = documents[1:]
            for i in document:
                if "/" not in i or len(i.split("/")) != 2:
                    print(f"无效的文档格式: {i}")
                    log(f"无效的文档格式: {i}", "warning", logfile_only=True)
                    document.remove(i)
                    continue
                if i.split("/")[1] == "annual_summary":
                    docments_path = os.path.join(ftfpath, i.split("/")[0], "年度总结.docx")
                    annual_summary_count += 1
                else:
                    docments_path = os.path.join(ftfpath, i.split("/")[0], f"{i.split('/')[1]}月.docx")
                    normal_count += 1
                if not os.path.exists(docments_path):
                    if i.split("/")[1] == "annual_summary":
                        print(f"未找到{i.split('/')[0]}年的年度总结")
                        log(f"未找到{i.split('/')[0]}年的年度总结", "warning", logfile_only=True)
                    else:
                        print(f"未找到{i}的事件记录文档")
                        log(f"未找到{i}的事件记录文档", "warning", logfile_only=True)
                    continue
                count = self._find(docments_path, keywords, count)
            self._check_month_span()
            if annual_summary_count > 0 and normal_count > 0:
                print(f"在{', '.join(document)}这{normal_count}个事件记录文档以及{annual_summary_count}个年度总结中共发现{count}个关键字词\n")
                log(f"在{', '.join(document)}这{normal_count}个事件记录文档以及{annual_summary_count}个年度总结中共发现{count}个关键字词", "info", logfile_only=True)
            elif annual_summary_count > 0 and normal_count == 0:
                print(f"在{', '.join(document)}这{annual_summary_count}个年度总结中共发现{count}个关键字词\n")
                log(f"在{', '.join(document)}这{annual_summary_count}个年度总结中共发现{count}个关键字词", "info", logfile_only=True)
            else:
                print(f"在{', '.join(document)}这{len(document)}个事件记录文档中共发现{count}个关键字词\n")
                log(f"在{', '.join(document)}这{len(document)}个事件记录文档中共发现{count}个关键字词", "info", logfile_only=True)
            log("", "info", logfile_only=True)
        else:
            print(self.do_find.__doc__)

    def complete_find(self, text: str, line: str, begidx: int, endidx: int) -> list[str]:
        if re.match(r"find [^.*$]+ in year \w*", line):
            return [i for i in self.YEARS if i.startswith(text)]
        if re.match(r"find [^.*$]+ in month \w*", line):
            return [i for i in [str(i) for i in range(1, 13)] if i.startswith(text)]
        if re.match(r"find [^.*$]+ in only \w*", line):
            parts = text.split("/")
            year_part = parts[0] if len(parts) > 0 else ""
            month_part = parts[1] if len(parts) > 1 else ""
            if "/" not in text:
                return [f"{y}/" for y in self.YEARS if y.startswith(year_part)]
            options = []
            for month in range(1, 13):
                month_str = str(month)
                if month_str.startswith(month_part):
                    options.append(f"{year_part}/{month_str}")
            if "annual_summary".startswith(month_part):
                options.append(f"{year_part}/annual_summary")
            return options
        if re.match(r"find [^.*$]+ in ", line) and len(line.split(" ")) == 4:
            return [i for i in ["*", "year", "month", "only"] if i.startswith(text)]
        if len(line.split(" ")) == 3:
            return ["in"]
        return []
    
    def do_open(self, args: str):
        """
        打开指定的文档。

        语法：open <document> [/?]
            document    指定的文档。
                        若文档格式为<year>/<month>，则表示打开事件记录文档，如“2023/1”表示2023年1月的事件记录文档。
                        若文档格式为<year>/annual_summary，则表示打开年度总结，如2023/annual_summary表示2023
                        年的年度总结。
            /?          显示此帮助文档。
        """
        if args.split(" ")[0] == "/?" or args == "":
            print(self.do_open.__doc__)
            return
        document = args
        if document.split("/")[1] == "annual_summary":
            docments_path = os.path.join(ftfpath, document.split("/")[0], "年度总结.docx")
        else:
            docments_path = os.path.join(ftfpath, document.split("/")[0], f"{document.split("/")[1]}月.docx")
        if not os.path.exists(docments_path):
            if document.split("/")[1] == "annual_summary":
                print(f"未找到{document.split('/')[0]}年的年度总结")
                log(f"未找到{document.split('/')[0]}年的年度总结", "warning", logfile_only=True)
            else:
                print(f"未找到{document}的事件记录文档")
                log(f"未找到{document}的事件记录文档", "warning", logfile_only=True)
            return
        os.system(f"start {docments_path}")
        if document.split("/")[1] == "annual_summary":
            log(f"打开了{document.split('/')[0]}年的年度总结", "info", logfile_only=True)
        else:
            log(f"打开了{document}的事件记录文档", "info", logfile_only=True)

    def complete_open(self, text: str, line: str, begidx: int, endidx: int) -> list[str]:
        parts = text.split("/")
        year_part = parts[0] if len(parts) > 0 else ""
        month_part = parts[1] if len(parts) > 1 else ""
        if "/" not in text:
            return [f"{i}/" for i in self.YEARS if i.startswith(year_part)]
        if month_part == "":
            months = [f"{year_part}/{i}" for i in [str(i) for i in range(1, 13)]]
            annual = [f"{year_part}/annual_summary"]
            return months + annual
        else:
            if month_part.startswith("a"):
                return [f"{year_part}/annual_summary"] if "annual_summary".startswith(month_part) else []
            else:
                return [f"{year_part}/{i}" for i in [str(i) for i in range(1, 13)] if i.startswith(month_part)]

    def do_count(self, args: str):
        """
        统计年度总结中的新事物（或重大事件）以及“时期”的数量。

        语法：count [N&M | normal_period | combined_period] in <year> [/?]
            N&M                 统计新事物（或重大事件）的数量。
            normal_period       统计常规“时期”的数量。
            combined_period     统计合称“时期”的数量。
            year                年度总结所属的年份。
            /?                  显示此帮助文档。
        """
        if args.split(" ")[0] == "/?" or args == "":
            print(self.do_count.__doc__)
            return
        if args.split(" ")[0] == "N&M":
            year = args.split(" ")[2]
            docments_path = os.path.join(ftfpath, year, "年度总结.docx")
            if not os.path.exists(docments_path):
                print(f"未找到{year}年的年度总结")
                log(f"未找到{year}年的年度总结", "warning", logfile_only=True)
                return
            doc = Document(docments_path)
            total_count = 0
            month_count = 0
            month_info = {}
            actual_records = 0
            actual_records_month = {}
            wrong_records_count = 0
            start_count = False
            for paragraph in doc.paragraphs:
                if start_count:
                    if paragraph.text == "":
                        break
                    if not paragraph.text[0].isdigit() and not paragraph.text.endswith("："):
                        total_count += 1
                        month_count += 1
                        month_info[month] = month_count
                    else:
                        month_count = 0
                        month = int(re.search(r"(\d+)月", paragraph.text).group(1))
                        actual_records_month[month] = int(re.search(r"共(\d+)个", paragraph.text).group(1)) if re.search(r"共(\d+)个", paragraph.text) else 0
                if not start_count and "新事物（或重大事件）" in paragraph.text:
                    start_count = True
                    actual_records = int(re.search(r"(\d+)个新事物（或重大事件）", paragraph.text).group(1)) if re.search(r"(\d+)个新事物（或重大事件）", paragraph.text) else 0
            print(f"在{year}年的年度总结中共发现{total_count}个新事物（或重大事件）")
            log(f"在{year}年的年度总结中共发现{total_count}个新事物（或重大事件）", "info", logfile_only=True)
            for k, v in month_info.items():
                if v == actual_records_month[k]:
                    print(f"{str(k)}月: {v}个")
                    log(f"{str(k)}月: {v}个", "info", logfile_only=True)
                else:
                    print(f"{str(k)}月: {v}个（年度总结中记录为{actual_records_month[k]}个，请更正）")
                    log(f"{str(k)}月: {v}个（年度总结中记录为{actual_records_month[k]}个，请更正）", "info", logfile_only=True)
                    wrong_records_count += 1
            if total_count == actual_records:
                print(f"在年度总结中共发现{total_count}个新事物（或重大事件），与实际记录一致")
                log(f"在年度总结中共发现{total_count}个新事物（或重大事件），与实际记录一致", "info", logfile_only=True)
            else:
                print(f"在年度总结中共发现{total_count}个新事物（或重大事件），但年度总结中记录为{actual_records}个，请更正")
                log(f"在年度总结中共发现{total_count}个新事物（或重大事件），但年度总结中记录为{actual_records}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            print(f"对{year}年新事物（或重大事件）的统计与检查已完成，共发现{wrong_records_count}项记录错误")
            log(f"对{year}年新事物（或重大事件）的统计与检查已完成，共发现{wrong_records_count}项记录错误", "info", logfile_only=True)
            print()
            log("", "info", logfile_only=True)
        elif args.split(" ")[0] == "normal_period":
            year = args.split(" ")[2]
            docments_path = os.path.join(ftfpath, year, "年度总结.docx")
            if not os.path.exists(docments_path):
                print(f"未找到{year}年的年度总结")
                log(f"未找到{year}年的年度总结", "warning", logfile_only=True)
                return
            doc = Document(docments_path)
            total_count = 0
            total_count_strong = 0
            total_count_weak = 0
            actual_total_count_strong = 0
            actual_total_count_weak = 0
            month_count = 0
            month_info = {}
            actual_records = 0
            wrong_records_count = 0
            start_count = False
            for paragraph in doc.paragraphs:
                if start_count:
                    if paragraph.text == "":
                        break
                    if not paragraph.text[0].isdigit() and not paragraph.text.endswith("："):
                        total_count += 1
                        month_count += 1
                        month_info[month] = month_count
                    else:
                        month_count = 0
                        month = int(re.search(r"(\d+)月", paragraph.text).group(1))
                    if "时期（强）" in paragraph.text:
                        total_count_strong += 1
                    if "时期（弱）" in paragraph.text:
                        total_count_weak += 1
                if not start_count and "常规“时期”" in paragraph.text:
                    start_count = True
                    actual_records = int(re.search(r"(\d+)个常规“时期”", paragraph.text).group(1)) if re.search(r"(\d+)个常规“时期”", paragraph.text) else 0
                    actual_total_count_strong = int(re.search(r"(\d+)个强“时期”", paragraph.text).group(1)) if re.search(r"(\d+)个强“时期”", paragraph.text) else 0
                    actual_total_count_weak = int(re.search(r"(\d+)个弱“时期”", paragraph.text).group(1)) if re.search(r"(\d+)个弱“时期”", paragraph.text) else 0
            print(f"在{year}年的年度总结中共发现{total_count}个常规“时期”")
            log(f"在{year}年的年度总结中共发现{total_count}个常规“时期”", "info", logfile_only=True)
            print(f"其中强“时期”有{total_count_strong}个，弱“时期”有{total_count_weak}个")
            log(f"其中强“时期”有{total_count_strong}个，弱“时期”有{total_count_weak}个", "info", logfile_only=True)
            for k, v in month_info.items():
                print(f"{str(k)}月: {v}个")
                log(f"{str(k)}月: {v}个", "info", logfile_only=True)
            if total_count == actual_records:
                print(f"在年度总结中共发现{total_count}个常规“时期”，与实际记录一致")
                log(f"在年度总结中共发现{total_count}个常规“时期”，与实际记录一致", "info", logfile_only=True)
            else:
                print(f"在年度总结中共发现{total_count}个常规“时期”，但年度总结中记录为{actual_records}个，请更正")
                log(f"在年度总结中共发现{total_count}个常规“时期”，但年度总结中记录为{actual_records}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            if total_count_strong == actual_total_count_strong:
                print(f"发现强“时期”{total_count_strong}个，与实际记录一致")
                log(f"发现强“时期”{total_count_strong}个，与实际记录一致", "info", logfile_only=True)
            else:
                print(f"发现强“时期”{total_count_strong}个，但年度总结中记录为{actual_total_count_strong}个，请更正")
                log(f"发现强“时期”{total_count_strong}个，但年度总结中记录为{actual_total_count_strong}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            if total_count_weak == actual_total_count_weak:
                print(f"发现弱“时期”{total_count_weak}个，与实际记录一致")
                log(f"发现弱“时期”{total_count_weak}个，与实际记录一致", "info", logfile_only=True)
            else:
                print(f"发现弱“时期”{total_count_weak}个，但年度总结中记录为{actual_total_count_weak}个，请更正")
                log(f"发现弱“时期”{total_count_weak}个，但年度总结中记录为{actual_total_count_weak}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            print(f"对{year}年常规“时期”的统计与检查已完成，共发现{wrong_records_count}项记录错误")
            log(f"对{year}年常规“时期”的统计与检查已完成，共发现{wrong_records_count}项记录错误", "info", logfile_only=True)
            print()
            log("", "info", logfile_only=True)
        elif args.split(" ")[0] == "combined_period":
            year = args.split(" ")[2]
            docments_path = os.path.join(ftfpath, year, "年度总结.docx")
            if not os.path.exists(docments_path):
                print(f"未找到{year}年的年度总结")
                log(f"未找到{year}年的年度总结", "warning", logfile_only=True)
                return
            doc = Document(docments_path)
            total_count = 0
            total_count_strong = 0
            total_count_weak = 0
            actual_total_count_strong = 0
            actual_total_count_weak = 0
            month_info = []
            actual_records = 0
            wrong_records_count = 0
            start_count = False
            for paragraph in doc.paragraphs:
                if start_count:
                    if paragraph.text == "":
                        break
                    total_count += 1
                    month_info.append(re.search(r"\d+~\d+月", paragraph.text).group(0))
                    if re.search(r"强.*“时期”", paragraph.text):
                        total_count_strong += 1
                    if re.search(r"弱.*“时期”", paragraph.text):
                        total_count_weak += 1
                if not start_count and "合称“时期”" in paragraph.text:
                    start_count = True
                    actual_records = int(re.search(r"(\d+)个合称“时期”", paragraph.text).group(1)) if re.search(r"(\d+)个合称“时期”", paragraph.text) else 0
                    actual_total_count_strong = int(re.search(r"(\d+)个强“时期”", paragraph.text).group(1)) if re.search(r"(\d+)个强“时期”", paragraph.text) else 0
                    actual_total_count_weak = int(re.search(r"(\d+)个弱“时期”", paragraph.text).group(1)) if re.search(r"(\d+)个弱“时期”", paragraph.text) else 0
            if total_count == actual_records:
                print(f"在{year}年的年度总结中共发现{total_count}个合称“时期”: {', '.join([i for i in month_info])}")
                log(f"在{year}年的年度总结中共发现{total_count}个合称“时期”: {', '.join([i for i in month_info])}", "info", logfile_only=True)
            else:
                print(f"在{year}年的年度总结中共发现{total_count}个合称“时期”，但年度总结中记录为{actual_records}个，请更正")
                log(f"在{year}年的年度总结中共发现{total_count}个合称“时期”，但年度总结中记录为{actual_records}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            if total_count_strong == actual_total_count_strong:
                print(f"发现强“时期”{total_count_strong}个，与实际记录一致")
                log(f"发现强“时期”{total_count_strong}个，与实际记录一致", "info", logfile_only=True)
            else:
                print(f"发现强“时期”{total_count_strong}个，但年度总结中记录为{actual_total_count_strong}个，请更正")
                log(f"发现强“时期”{total_count_strong}个，但年度总结中记录为{actual_total_count_strong}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            if total_count_weak == actual_total_count_weak:
                print(f"发现弱“时期”{total_count_weak}个，与实际记录一致")
                log(f"发现弱“时期”{total_count_weak}个，与实际记录一致", "info", logfile_only=True)
            else:
                print(f"发现弱“时期”{total_count_weak}个，但年度总结中记录为{actual_total_count_weak}个，请更正")
                log(f"发现弱“时期”{total_count_weak}个，但年度总结中记录为{actual_total_count_weak}个，请更正", "info", logfile_only=True)
                wrong_records_count += 1
            print(f"对{year}年合称“时期”的统计与检查已完成，共发现{wrong_records_count}项记录错误")
            log(f"对{year}年合称“时期”的统计与检查已完成，共发现{wrong_records_count}项记录错误", "info", logfile_only=True)
            print()
            log("", "info", logfile_only=True)
        else:
            print(self.do_count.__doc__)

    def complete_count(self, text: str, line: str, begidx: int, endidx: int) -> list[str]:
        if re.match(r"count (N&M|normal_period|combined_period) in \w*", line):
            return [i for i in self.YEARS if i.startswith(text)]
        if len(line.split(" ")) > 2:
            return ["in"]
        if line.startswith("count "):
            return [i for i in ["N&M", "normal_period", "combined_period"] if i.startswith(text)]
        return []
    
    def _check_annual_summary(self, paragraphs: List[Paragraph], year: str):
        line = 0
        normal_period_start = False
        combined_period_start = False
        irregularity_count = 0
        for paragraph in paragraphs:
            line += 1
            if line == 1 and paragraph.text != "年度总结":
                print(f"{year}年年度总结第{line}个段落: 标题不规范（应为“年度总结”），请更正")
                log(f"{year}年年度总结第{line}个段落: 标题不规范（应为“年度总结”），请更正", "info", logfile_only=True)
                irregularity_count += 1
            if line == 2 and paragraph.text != "朝花已经绽放，是时候将它拾起":
                print(f"{year}年年度总结第{line}个段落: 应为“朝花已经绽放，是时候将它拾起”，请更正")
                log(f"{year}年年度总结第{line}个段落: 应为“朝花已经绽放，是时候将它拾起”，请更正", "info", logfile_only=True)
                irregularity_count += 1
            year_match = re.search(r"(\d+)年共有(\d+)个新事物（或重大事件）", paragraph.text)
            if year_match and year_match.group(1) != year:
                print(f"{year}年年度总结第{line}个段落: 年份不匹配，请更正")
                log(f"{year}年年度总结第{line}个段落: 年份不匹配，请更正", "info", logfile_only=True)
                irregularity_count += 1
            if paragraph.text != "" and paragraph.text[0].isdigit() and paragraph.text[1] == "月":
                if int(paragraph.text[0]) < 1 or int(paragraph.text[0]) > 12:
                    print(f"{year}年年度总结第{line}个段落: 月份异常（月份应为1到12的数字），请更正")
                    log(f"{year}年年度总结第{line}个段落: 月份异常（月份应为1到12的数字），请更正", "info", logfile_only=True)
                    irregularity_count += 1
                if not paragraph.text.endswith("："):
                    print(f"{year}年年度总结第{line}个段落: 缺少冒号，请更正")
                    log(f"{year}年年度总结第{line}个段落: 缺少冒号，请更正", "info", logfile_only=True)
                    irregularity_count += 1
                if not re.search(r"共\d+个", paragraph.text) and not normal_period_start and not combined_period_start:
                    print(f"{year}年年度总结第{line}个段落: 缺少新事物（或重大事件）的数量统计（应包含“共X个”字样），请更正")
                    log(f"{year}年年度总结第{line}个段落: 缺少新事物（或重大事件）的数量统计（应包含“共X个”字样），请更正", "info", logfile_only=True)
                    irregularity_count += 1
            if paragraph.text.startswith("其中，共有"):
                normal_period_start = True
                if paragraphs[line - 2].text != "":
                    print(f"{year}年年度总结第{line}个段落: 未空行，请更正")
                    log(f"{year}年年度总结第{line}个段落: 未空行，请更正", "info", logfile_only=True)
                    irregularity_count += 1
                if not re.search(r"(\d+)个常规“时期”", paragraph.text):
                    print(f"{year}年年度总结第{line}个段落: 未统计常规“时期”数量，请更正")
                    log(f"{year}年年度总结第{line}个段落: 未统计常规“时期”数量，请更正", "info", logfile_only=True)
                    irregularity_count += 1
            if paragraph.text != "" and paragraph.text[0].isdigit() and paragraph.text[1] == "月" and "称为" not in paragraph.text and normal_period_start:
                print(f"{year}年年度总结第{line}个段落: 缺少月份命名，请更正")
                log(f"{year}年年度总结第{line}个段落: 缺少月份命名，请更正", "info", logfile_only=True)
                irregularity_count += 1
            if paragraph.text.startswith("以及"):
                combined_period_start = True
                if paragraphs[line - 2].text != "":
                    print(f"{year}年年度总结第{line}个段落: 未空行，请更正")
                    log(f"{year}年年度总结第{line}个段落: 未空行，请更正", "info", logfile_only=True)
                    irregularity_count += 1
                if not re.search(r"(\d+)个合称“时期”", paragraph.text):
                    print(f"{year}年年度总结第{line}个段落: 未统计合称“时期”数量，请更正")
                    log(f"{year}年年度总结第{line}个段落: 未统计合称“时期”数量，请更正", "info", logfile_only=True)
                    irregularity_count += 1
            if paragraph.text.startswith("年主题曲") and paragraphs[line - 2].text != "":
                print(f"{year}年年度总结第{line}个段落: 未空行，请更正")
                log(f"{year}年年度总结第{line}个段落: 未空行，请更正", "info", logfile_only=True)
                irregularity_count += 1
        if "年主题曲：" not in str([p.text for p in paragraphs]):
            print(f"{year}年年度总结: 缺少年主题曲部分（若没有也请标记为“无”），请更正")
            log(f"{year}年年度总结: 缺少年主题曲部分（若没有也请标记为“无”），请更正", "info", logfile_only=True)
            irregularity_count += 1
        if "年文章：" not in str([p.text for p in paragraphs]):
            print(f"{year}年年度总结: 缺少年文章部分（若没有也请标记为“无”），请更正")
            log(f"{year}年年度总结: 缺少年文章部分（若没有也请标记为“无”），请更正", "info", logfile_only=True)
            irregularity_count += 1
        possible_yearly_assessment = [p.text for p in paragraphs][-1]
        if possible_yearly_assessment.startswith("年度评估："):
            yearly_assessment = possible_yearly_assessment.split("：")[1]
            if yearly_assessment not in self.YEARLY_JUDGMENT:
                print(f"{year}年年度总结: 年度评估内容无效，请更正")
                log(f"{year}年年度总结: 年度评估内容无效，请更正", "info", logfile_only=True)
                irregularity_count += 1
        else:
            print(f"{year}年年度总结: 缺少年度评估，请更正")
            log(f"{year}年年度总结: 缺少年度评估，请更正", "info", logfile_only=True)
            irregularity_count += 1
        return irregularity_count

    def _check_event_record(self, paragraphs: List[Paragraph], year: str, month: str):
        irregularity_count = 0
        parts = []
        possible_monthly_assessment = []
        for paragraph in paragraphs:
            if paragraph.text == "":
                parts.append(possible_monthly_assessment)
                possible_monthly_assessment = []
                continue
            possible_monthly_assessment.append(paragraph.text)
        for part in parts:
            if not part:
                continue
            if not re.search(r"(\d+)年(\d+)月", part[0]):
                print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 未记录时间，请更正")
                log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 未记录时间，请更正", "info", logfile_only=True)
                irregularity_count += 1
            if "周度：" in part[0]:
                possible_weekly_assessment = part[-1]
                if possible_weekly_assessment.startswith("周度评估："):
                    weekly_assessment = possible_weekly_assessment.split("：")[1]
                    if weekly_assessment not in self.WEEKLY_JUDGMENT:
                        print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 周度评估内容无效，请更正")
                        log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 周度评估内容无效，请更正", "info", logfile_only=True)
                        irregularity_count += 1
                else:
                    print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 缺少周度评估，请更正")
                    log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 缺少周度评估，请更正", "info", logfile_only=True)
                    irregularity_count += 1
            if not "正面情感评估：" in str(part):
                print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 缺少正面情感评估，请更正")
                log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 缺少正面情感评估，请更正", "info", logfile_only=True)
                irregularity_count += 1
            else:
                for line in part:
                    if line.startswith("正面情感评估："):
                        if "，" in line:
                            positive_level = line.split("：")[1].split("，")[0]
                            positive_assessment = line.split("：")[1].split("，")[1]
                        else:
                            positive_level = line.split("：")[1]
                            positive_assessment = None
                        if positive_level not in self.POSITIVE_LEVELS:
                            print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 正面情感类型等级无效，请更正")
                            log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 正面情感类型等级无效，请更正", "info", logfile_only=True)
                            irregularity_count += 1
                        if positive_assessment and positive_assessment not in self.POSITIVE_ASSESS:
                            print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 积极情感评估等级无效，请更正")
                            log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 积极情感评估等级无效，请更正", "info", logfile_only=True)
                            irregularity_count += 1
            if not "负面情感评估：" in str(part):
                print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 缺少负面情感评估，请更正")
                log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 缺少负面情感评估，请更正", "info", logfile_only=True)
                irregularity_count += 1
            else:
                for line in part:
                    if line.startswith("负面情感评估："):
                        if "，" in line:
                            negative_level = line.split("：")[1].split("，")[0]
                            negative_assessment = line.split("：")[1].split("，")[1]
                        else:
                            negative_level = line.split("：")[1]
                            negative_assessment = None
                        if negative_level not in self.NEGATIVE_LEVELS:
                            print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 负面情感类型等级无效，请更正")
                            log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 负面情感类型等级无效，请更正", "info", logfile_only=True)
                            irregularity_count += 1
                        if negative_assessment and negative_assessment not in self.NEGATIVE_ASSESS:
                            print(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 消极情感评估等级无效，请更正")
                            log(f"{year}年{month}月事件记录文档第{parts.index(part) + 1}周: 消极情感评估等级无效，请更正", "info", logfile_only=True)
                            irregularity_count += 1
        if possible_monthly_assessment and possible_monthly_assessment[0].startswith("月度评估："):
            monthly_assessment = possible_monthly_assessment[0].split("：")[1]
            if monthly_assessment not in self.MONTHLY_JUDGMENT:
                print(f"{year}年{month}月事件记录文档: 月度评估内容无效，请更正")
                log(f"{year}年{month}月事件记录文档: 月度评估内容无效，请更正", "info", logfile_only=True)
                irregularity_count += 1
        else:
            print(f"{year}年{month}月事件记录文档: 缺少月度评估，但也可能是该月尚未结束，请留意")
            log(f"{year}年{month}月事件记录文档: 缺少月度评估，但也可能是该月尚未结束，请留意", "info", logfile_only=True)
            irregularity_count += 1
        return irregularity_count

    def _check(self, document: str):
        filename = os.path.basename(document)
        if filename.startswith("~$"):
            return 0
        doc = Document(document)
        year = os.path.basename(os.path.dirname(document))
        irregularity_count = 0
        if filename == "年度总结.docx":
            irregularity_count += self._check_annual_summary(doc.paragraphs, year)
        if filename.endswith("月.docx"):
            irregularity_count += self._check_event_record(doc.paragraphs, year, filename.split("月")[0])
        return irregularity_count

    def do_check(self, args: str):
        """
        检查事件记录文档及年度总结的格式规范性。

        语法：check <years> [/?]
            years   指定要检查的年份范围，可填多个，使用空格分隔。
                    也可使用“~”表示范围，如“2023~2026”表示检查2023年到2026年（包含2023年和2026年）的文档。
                    键入“*”表示检查所有文档。
            /?      显示此帮助文档。
        
        注：由于2023年的文档处于早期阶段，并未统一格式，因此不会检查2023年的文档。
        """
        if args.split(" ")[0] == "/?" or args == "":
            print(self.do_check.__doc__)
            return
        years = args.split(" ")
        if years[0] == "*":
            years = self.YEARS[1:]
        else:
            deduplicated_year = set()
            for year in years:
                if "~" in year:
                    start_year, end_year = year.split("~")
                    if not start_year.isdigit() or not end_year.isdigit():
                        print(f"无效的年份范围: {year}")
                        log(f"无效的年份范围: {year}", "warning", logfile_only=True)
                        years.remove(year)
                        continue
                    start_year, end_year = int(start_year), int(end_year)
                    if start_year > end_year:
                        print(f"无效的年份范围: {year}（起始年份大于结束年份）")
                        log(f"无效的年份范围: {year}（起始年份大于结束年份）", "warning", logfile_only=True)
                        years.remove(year)
                        continue
                    deduplicated_year.update(str(i) for i in range(start_year, end_year + 1))
                else:
                    if not year.isdigit():
                        print(f"无效的年份: {year}")
                        log(f"无效的年份: {year}", "warning", logfile_only=True)
                        continue
                    deduplicated_year.add(year)
            years = natsorted(deduplicated_year)
            if "2023" in years:
                print("通常不会检查2023年的文档，具体原因参见帮助文档。")
                log("通常不会检查2023年的文档，具体原因参见帮助文档。", "info", logfile_only=True)
                years.remove("2023")
            if len(years) == 0:
                print("未指定有效的年份，检查已取消。")
                log("未指定有效的年份，检查已取消。", "warning", logfile_only=True)
                return
        irregularity_count = 0
        for year in years:
            docments_path = os.path.join(ftfpath, year)
            for document in natsorted(glob(f"{docments_path}\\*.docx")):
                irregularity_count += self._check(document)
        print(f"对{', '.join(years)}年文档的格式检查已完成，共发现{irregularity_count}项不规范处")
        log(f"对{', '.join(years)}年文档的格式检查已完成，共发现{irregularity_count}项不规范处", "info", logfile_only=True)
        print("注意: 检查仅针对格式规范性，对于统计等数值部分的正确性并不进行判断")
        log("注意: 检查仅针对格式规范性，对于统计等数值部分的正确性并不进行判断", "info", logfile_only=True)

    def complete_check(self, text: str, line: str, begidx: int, endidx: int) -> list[str]:
        if line.startswith("check "):
            return [i for i in self.YEARS if i.startswith(text)]
        return []

    def do_calculate(self, args: str):
        """
        计算“时期”的强度。

        语法：calculate [strong <count>] [weak <count>] [scattered <coefficient>] start <time> end <time>/now [/?]
            strong <count>              强定位物的数量。
            weak <count>                弱定位物的数量。
            scattered <coefficient>     若为分散性“时期”，指定的分散系数，否则默认为1。
            start <month>               “时期”记录起算时间，格式为“年份/月份”，例如“2025/1”。
            end <month>/now             指定结束时间，或键入“now”以选择当前时间。
            /?                          显示此帮助文档。
        
        注：以上输入参数的顺序可以任意调整，但起算时间必须早于结束时间。
        """
        if args.split(" ")[0] == "/?" or args == "":
            print(self.do_calculate.__doc__)
            return
        parts = args.split(" ")
        if "start" not in parts or "end" not in parts:
            print(self.do_calculate.__doc__)
            return
        start_index = parts.index("start")
        end_index = parts.index("end")
        if start_index > end_index:
            print(self.do_calculate.__doc__)
            return
        start_time = parts[start_index + 1]
        end_time = parts[end_index + 1]
        strong_count = int(parts[parts.index("strong") + 1]) if "strong" in parts else 0
        weak_count = int(parts[parts.index("weak") + 1]) if "weak" in parts else 0
        scattered_coefficient = int(parts[parts.index("scattered") + 1]) if "scattered" in parts else 1
        start_year, start_month = map(int, start_time.split("/"))
        if end_time == "now":
            now = datetime.datetime.now()
            end_year, end_month = now.year, now.month
            end_time = f"{end_year}/{end_month}"
        else:
            end_year, end_month = map(int, end_time.split("/"))
        if (start_year > end_year) or (start_year == end_year and start_month > end_month):
            print("结束时间必须晚于起算时间")
            log("结束时间必须晚于起算时间", "warning", logfile_only=True)
            return
        interval_months = (end_year - start_year) * 12 + (end_month - start_month)
        total_items = strong_count + weak_count
        total_strength = strong_count * 2 + weak_count * 1
        intensity = ((total_items * total_strength) / scattered_coefficient) * math.exp(-0.1 * interval_months)
        print(f"从{start_time}到{end_time}，间隔{interval_months}个月")
        print(f"强定位物数量: {strong_count}")
        print(f"弱定位物数量: {weak_count}")
        print(f"分散系数: {scattered_coefficient}")
        print(f"计算得到的“时期”强度为: {intensity:.2f}夕")
        log(f"计算得到的“时期”强度为: {intensity:.2f}夕", "info", logfile_only=True)

    def complete_calculate(self, text: str, line: str, begidx: int, endidx: int) -> list[str]:
        options = ["strong", "weak", "scattered", "start", "end"]
        parts = line.split(" ")
        if len(parts) > 1:
            if parts[-1] in options:
                return []
            elif parts[-2] == "start":
                if "/" in text:
                    parts = text.split("/")
                    if len(parts) == 2:
                        year_part, month_part = parts
                        if year_part in self.YEARS or any(y.startswith(year_part) for y in self.YEARS):
                            months = [str(i) for i in range(1, 13)]
                            completions = [f"{year_part}/{m}" for m in months if m.startswith(month_part)]
                            return completions
                    return []
                else:
                    year_completions = [y for y in self.YEARS if y.startswith(text)]
                    return [f"{y}/" for y in year_completions]
            elif parts[-2] == "end":
                if "/" in text:
                    parts = text.split("/")
                    if len(parts) == 2:
                        year_part, month_part = parts
                        if year_part in self.YEARS or any(y.startswith(year_part) for y in self.YEARS):
                            months = [str(i) for i in range(1, 13)]
                            completions = [f"{year_part}/{m}" for m in months if m.startswith(month_part)]
                            return completions
                    return []
                else:
                    if text.startswith("n"):
                        return ["now"]
                    year_completions = [f"{y}/" for y in self.YEARS if y.startswith(text)] + ["now"]
                    return [y for y in year_completions if y.startswith(text)]
            else:
                return [i for i in options if i.startswith(parts[-1])]
        else:
            return options

def help_ftf():
    FTFCmd.cipher = getpass.getpass("《朝花夕拾协议》的宗旨是？")
    if not bcrypt.checkpw(FTFCmd.cipher.encode(), b'$2b$12$aI7vpFUGDjIk0wKLXYSTE./UHK3TDZlH9/XMF7Jbf5dLx5pfEntUi'):
        return
    for i in range(2):
        os.system("color 0c")
        sleep(0.1)
        os.system("color 0a")
        sleep(0.1)
    sleep(2)
    raise AdminMode()

class FTFAdminCmd(FTFCmd):
    prompt = "[FTF ADMIN] "

    def __init__(self, completekey = "tab", stdin = None, stdout = None):
        super().__init__(completekey, stdin, stdout)
        self.COLOR = Fore.LIGHTRED_EX

    def onecmd(self, line: str) -> bool:
        if line == "" or line.isspace():
            return
        c = choice("YN", Fore.LIGHTYELLOW_EX + "协议创始人，请牢记您的命令一旦执行便无法撤回，您确定要执行此命令吗")
        if c == 2:
            return
        os.system("color 0c")
        return super().onecmd(line)
    
    def do_exit(self, args: str):
        """
        退出《朝花夕拾协议》命令行。

        语法：exit [/?]
            无参数  退出《朝花夕拾协议》命令行。
            /?      显示此帮助文档。
        """
        if args.split(" ")[0] == "/?":
            print(self.do_exit.__doc__)
            return
        os.system("color 0a")
        raise CommandLineExit()
    
    def do_dellog(self, args: str):
        """
        删除日志文件。

        语法：dellog <logname/*> [/?]
            logname     指定删除的日志文件的名称。
            *           表示所有日志文件。
            /?          显示此帮助文档。
        """
        if args.split(" ")[0] == "/?" or args == "":
            print(self.do_dellog.__doc__)
            return
        logname = args
        if logname == "*":
            for i in glob("logs\\*.log"):
                try:
                    os.remove(i)
                except PermissionError:
                    pass
                print(f"已删除{i}")
                log(f"已删除{i}", "info", logfile_only=True)
        else:
            if not os.path.exists(f"logs\\{logname}.log"):
                print(f"未找到{logname}的日志文件")
                log(f"未找到{logname}的日志文件", "warning", logfile_only=True)
                return
            try:
                os.remove(f"logs\\{logname}.log")
            except PermissionError:
                print(f"无法删除{logname}.log，请确保文件未被占用")
                log(f"无法删除{logname}.log，请确保文件未被占用", "warning", logfile_only=True)
                return
            print(f"已删除{logname}.log")
            log(f"已删除{logname}.log", "info", logfile_only=True)

    def complete_dellog(self, text: str, line: str, begidx: int, endidx: int) -> list[str]:
        if line.startswith("dellog "):
            return [i.replace("logs\\", "").replace(".log", "") for i in glob("logs\\*.log") if i.replace("logs\\", "").replace(".log", "").startswith(text)]
        return []

def help_ftf_admin():
    print(Fore.LIGHTRED_EX + "您已处在协议创始人权限下")
